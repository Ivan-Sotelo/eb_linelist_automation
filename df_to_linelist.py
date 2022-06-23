import docx
from docx import Document
import pandas as pd
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches
from docx.shared import Pt
import numpy as np
from datetime import date
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from datetime import datetime
from docx.dml.color import ColorFormat
from docx.shared import RGBColor

def make_paragraph(document, font_name, font_size, alignment, text, bold):
    p = document.add_paragraph()
    title = p.add_run(text)
    title.bold = bold
    title.font.size = Pt(font_size)
    title.font.name = font_name
    p.alignment = alignment

def add_line_break(document):
    p = document.add_paragraph()
    run = p.add_run()
    run.add_break(docx.text.run.WD_BREAK.LINE)

def generate_table(df,table):
    for j in range(df.shape[-1]):
        table.cell(0,j).text = df.columns[j]
    
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            table.cell(i+1,j).text = str(df.values[i,j])

def set_cell_bg_color(tc, hex_color):
    """
    set background shading for Header Rows
    """
    tblCellProperties = tc._element.tcPr
    clShading = OxmlElement('w:shd')
    clShading.set(qn('w:fill'), hex_color)  # Hex of Dark Blue Shade {R:0x00, G:0x51, B:0x9E}
    tblCellProperties.append(clShading)

def set_cell_width(table, column, width):
   for cell in table.columns[column].cells:
       cell.width = Inches(width)

def set_alignment(table, alignment):
    # center the text in the cells
    if alignment == 'center':
        alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    for col in table.columns:
        for cell in col.cells:
            cell.paragraphs[0].alignment = alignment

def set_alignment_column(table, alignment, column_number):
    if alignment == 'center':
        alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    if alignment == 'left':
        alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
    for cell in table.columns[column_number].cells:
            cell.paragraphs[0].alignment = alignment

def set_font_size_column(table, column_number, size):
    for cell in table.columns[column_number].cells:
            paragraph = cell.paragraphs[0]
            font = paragraph.runs[0].font
            font.size = Pt(size)

def set_font_size(table, size):
    # set the font size of the text each cell
    for row in table.rows:
        for cell in row.cells:
            paragraph = cell.paragraphs[0]
            font = paragraph.runs[0].font
            font.size = Pt(size)

def set_header_bold(table,bold):
    #make a header text bold
    for cell in table.rows[0].cells:
            paragraph = cell.paragraphs[0]
            bold_status = paragraph.runs[0]
            bold_status.bold = bold

def add_cov22(linelist_df): 
    #add 'COV22-' to entries in RITM LAB ID column
    for index, row in linelist_df.iterrows():
        linelist_df.at[index, "RITM Lab ID"] = 'COV22-' + row['RITM Lab ID']
    return linelist_df

def make_outcome_column(linelist_df):
    for index, row in linelist_df.iterrows():
        if pd.isnull(linelist_df.at[index, "LINEAGE"]) == False:
             linelist_df.at[index, "OUTCOME"] = 'Lineage Assigned'
        else:
             linelist_df.at[index, "OUTCOME"] = 'Lineage Unassigned'
    return linelist_df

def combine_lineage(linelist_df):
    for index, row in linelist_df.iterrows():
        if pd.isnull(linelist_df.at[index, "LINEAGE"]) == False:                                            #processes lineage if value is not NAN
             linelist_df.at[index, "LINEAGE"] = row['LINEAGE'] + ' (' + row['lineage'] + ')'
        else:
            linelist_df.at[index, "LINEAGE"] = ''
    return linelist_df

def sort_and_manipulate_date_columns():
    #sorting and manipulating date_receieved column to become datetime objects
    date_list = linelist_df.date_received


    date_list = [datetime.strptime(datetime_str, '%m/%d/%Y') for datetime_str in date_list] #Y-M-D for date_recieved

    date_list = [datetime.strftime(datetime_str, '%B %d %Y') for datetime_str in date_list]


    linelist_df['date_received'] = date_list
    linelist_df['date_received'] =  pd.to_datetime(linelist_df['date_received'], format='%B %d %Y')  #makes column into datetime.datetime
    

    date_list = list(set(date_list))            # gets all unique dates and stores them in list
    

    date_list.sort(key = lambda date: datetime.strptime(date, '%B %d %Y'))      #sorts all dates
    

    date_list = [date_list[x:x+2] for x in range(0, len(date_list),2)]          #subset list of dates for every two date entries
    for dates in date_list:
        dates.sort(key = lambda date: datetime.strptime(date, '%B %d %Y'))
    

    #manipulating date of collection column to make them datetime objects
    date_collected = linelist_df['DATE OF COLLECTION']

    date_collected = [datetime.strptime(datetime_str, '%m/%d/%Y') for datetime_str in date_collected]   #M/D/Y for date_collected

    date_collected = [datetime.strftime(datetime_str, '%B %d %Y') for datetime_str in date_collected]

    linelist_df['DATE OF COLLECTION'] = date_collected
    linelist_df['DATE OF COLLECTION'] =  pd.to_datetime(linelist_df['DATE OF COLLECTION'], format='%B %d %Y')

    

    return date_list

def initialize_appendix_df():
    appendix_df = pd.DataFrame(columns=['Date specimens received at RITM'])
    date_received = []
    collection_period = []
    collection_period_subset = []
    num_samples_received = []
    rejected_samples = []
    num_rejected_samples = []
    num_lineage_assigned = []
    num_poor_genome = []
    lineage_list = []
    unique_lineage_list = []
    key_points = []

def subset_and_generate_linelist_table():
    #subset the dataframe every two dates
    for dates in date_list:
    
        header = document.add_table(rows=1, cols=1)
        header.style = 'Table Grid'

        if len(dates) == 2:
            date1 = dates[0]
            date1 = datetime.strptime(date1, '%B %d %Y')
            date2 = dates[1]
            date2 = datetime.strptime(date2, '%B %d %Y')
            
            
            linelist_subset = linelist_df[(linelist_df['date_received'] >= date1) & (linelist_df['date_received'] <= date2)]
            date1_str = date1.strftime("%B %d %Y")
            date2_str = date2.strftime("%B %d %Y")
            date_received.append(f'{date1_str} - {date2_str}')
                                                                                           #appends list to get date received range

            collection_period_subset = list(set(linelist_subset['DATE OF COLLECTION']))
            min_collect = min(collection_period_subset)
            min_collect = min_collect.strftime("%B %d %Y")
            max_collect = max(collection_period_subset)
            max_collect = max_collect.strftime("%B %d %Y")

            collection_period_subset = [date_str.strftime("%B %d %Y") for date_str in collection_period_subset]

                                                                                                                                #convert dates in dates collected to readable format
            linelist_subset['DATE OF COLLECTION'] = linelist_subset['DATE OF COLLECTION'].dt.strftime('%m/%d/%Y')
        
            if len(collection_period_subset) != 1:
                collection_period.append(min_collect + ' - ' + max_collect)                    #Appends list to get collection date range for this subset
            else:
                collection_period.append(min_collect)

            num_samples_received.append(linelist_subset.shape[0])

            rejected_samples = linelist_subset[linelist_subset['genome_coverage'].isnull()]                                     
            num_rejected_samples.append(rejected_samples.shape[0])                                                               # number of samples where there is no genome coverage

            lineage_assigned = linelist_subset[linelist_subset['OUTCOME'] == 'Lineage Assigned']
            num_lineage_assigned.append(lineage_assigned.shape[0])                                                               #number of samples where lineage was assigned

            poor_genome = linelist_subset[linelist_subset['genome_coverage'] < 70]
            num_poor_genome.append(poor_genome.shape[0])                                                                         # number of samples where genome coverage is less than 70

            lineage_list = list(linelist_subset['LINEAGE'][linelist_subset['LINEAGE'] != ''])
            unique_lineage_list = list(set(lineage_list))
            num_lineage_dict = {i:lineage_list.count(i) for i in unique_lineage_list}                                            # make a dictionary containing the unique lineages and their number of occurance

            key_point_string = f'· {round(((lineage_assigned.shape[0]/linelist_subset.shape[0])*100),2)}% ({lineage_assigned.shape[0]}/{linelist_subset.shape[0]}) successfully assigned with lineage \n'

            for key in num_lineage_dict.keys():
                key_point_string = key_point_string + f'· {round(((num_lineage_dict[key]/lineage_assigned.shape[0])*100),2)}% ({num_lineage_dict[key]}/{lineage_assigned.shape[0]}) {key} \n'


            header.cell(0,0).text = f'SAMPLES ENDORSED BY EPIDEMIOLOGY BUREAU (SAMPLES RECEIVED - {date1_str} TO {date2_str})'
        else:
            date1 = dates[0]
            linelist_subset = linelist_df[(linelist_df['date_received'] >= date1)]
        
            date_received.append(f'{date1}')
                                                                                                     #appends list to get date received range

            collection_period_subset = list(set(linelist_subset['DATE OF COLLECTION']))
            min_collect = min(collection_period_subset)
            min_collect = min_collect.strftime("%B %d %Y")
            max_collect = max(collection_period_subset)
            max_collect = max_collect.strftime("%B %d %Y")

            collection_period_subset = [date_str.strftime("%B %d %Y") for date_str in collection_period_subset]

                                                                                                                                #convert dates in dates collected to readable format
            linelist_subset['DATE OF COLLECTION'] = linelist_subset['DATE OF COLLECTION'].dt.strftime('%m/%d/%Y')

            if len(collection_period_subset) != 1:
                collection_period.append(min_collect + ' - ' + max_collect)                    #Appends list to get collection date range for this subset
            else:
                collection_period.append(min_collect)

            num_samples_received.append(linelist_subset.shape[0])                                                               #number of samples in this receiving period

            rejected_samples = linelist_subset[linelist_subset['genome_coverage'].isnull()]
            num_rejected_samples.append(rejected_samples.shape[0])                                                              # number of samples where there is no genome coverage

            lineage_assigned = linelist_subset[linelist_subset['OUTCOME'] == 'Lineage Assigned']
            num_lineage_assigned.append(lineage_assigned.shape[0])                                                               #number of samples where lineage was assigned

            poor_genome = linelist_subset[linelist_subset['genome_coverage'] < 70]
            num_poor_genome.append(poor_genome.shape[0])                                                                       # number of samples where genome coverage is less than 70

            lineage_list = list(linelist_subset['LINEAGE'][linelist_subset['LINEAGE'] != ''])
            unique_lineage_list = list(set(lineage_list))
            num_lineage_dict = {i:lineage_list.count(i) for i in unique_lineage_list}                                           # make a dictionary containing the unique lineages and their number of occurance
    
            key_point_string = f'· {round(((lineage_assigned.shape[0]/linelist_subset.shape[0])*100),2)}% ({lineage_assigned.shape[0]}/{linelist_subset.shape[0]}) successfully assigned with lineage \n'

            for key in num_lineage_dict.keys():
                key_point_string = key_point_string + f'· {round(((num_lineage_dict[key]/lineage_assigned.shape[0])*100),2)}% ({num_lineage_dict[key]}/{lineage_assigned.shape[0]}) {key} \n'

        
        
            header.cell(0,0).text = f'SAMPLES ENDORSED BY EPIDEMIOLOGY BUREAU (SAMPLES RECEIVED - {date1})'
    
    
        linelist_subset = linelist_subset[['NO', 'RITM Lab ID', 'UIC', 'AGE', 'SEX', 'PATIENT ADDRESS', 'LINELIST REGION', 'SENDING FACILITY', 'DATE OF COLLECTION', 'OUTCOME', 'LINEAGE','REMARKS']]
        linelist_subset['NO'] = np.arange(len(linelist_subset)) + 1

    


        set_cell_bg_color(header.cell(0,0), '084C22')
        set_font_size(header, 12)
        set_header_bold(header, True)
    
        eb_table = document.add_table(rows=linelist_subset.shape[0]+1, cols=linelist_subset.shape[1])
        eb_table.style = 'Table Grid'
        generate_table(linelist_subset, eb_table)

        set_cell_width(eb_table, 0, 0.4)   #NO
        set_cell_width(eb_table, 1, 1.37)  #RITM Lab ID
        set_cell_width(eb_table, 2, 1.08)  #UIC
        set_cell_width(eb_table, 3, 0.49)  #AGE
        set_cell_width(eb_table, 4, 0.4)   #SEX
        set_cell_width(eb_table, 5, 1.1)   #PATIENT ADDRESS
        set_cell_width(eb_table, 6, 0.7)   #LINELIST REGION
        set_cell_width(eb_table, 7, 1.65)  #SENDING FACILITY
        set_cell_width(eb_table, 8, 0.99)  #DATA OF COLLECTION
        set_cell_width(eb_table, 9, 0.86)  #OUTCOME
        set_cell_width(eb_table, 10, 1.97) #LINEAGE
        set_cell_width(eb_table, 11, 1.97) #REMARKS

        set_alignment(eb_table, 'center')

        set_font_size(eb_table, 11)
    
        set_header_bold(eb_table, True)

        set_cell_height(eb_table)
    
        document.add_page_break()

  
        key_points.append(key_point_string)

def populate_appendix_df():
    appendix_df['Date specimens received at RITM'] = date_received
    appendix_df['Date of Sequencing results release'] = ''
    appendix_df['Sample source'] = 'Targeted samples (endorsed by EB)'
    appendix_df['Period of specimen collection'] = collection_period
    appendix_df['N Received'] = num_samples_received
    appendix_df['Rejected'] = num_rejected_samples
    appendix_df['Ongoing Investigation'] = ''
    appendix_df['Lineage assigned'] = num_lineage_assigned
    appendix_df['Poor genome coverage'] = num_poor_genome
    appendix_df['Key points'] = key_points

def color_header(table, hex_color):
    #color the header row for table 1
    cell_count = 0
    for cell in table.rows[0].cells:
        set_cell_bg_color(table.cell(0,cell_count), hex_color)
        
        cell_count +=1

def set_table_1_params():
    #set the cell width and font size for table 1
    set_cell_width(table1, 0, 1.5)   #Date specimens received at RITM
    set_cell_width(table1, 1, 1.42)  #Date of Sequencing results release
    set_cell_width(table1, 2, 1.72)  #Sample source
    set_cell_width(table1, 3, 1.34)  #Period of specimen collection
    set_cell_width(table1, 4, 0.62)   #N Received
    set_cell_width(table1, 5, 0.65)   #Rejected
    set_cell_width(table1, 6, 0.82)   #Ongoing Investigation
    set_cell_width(table1, 7, 0.84)  #Lineage assigned
    set_cell_width(table1, 8, 0.82)  #Poor genome coverage
    set_cell_width(table1, 9, 3.26)  #Key points

    set_alignment(table1, 'center')

    set_font_size(table1, 9)

    set_alignment_column(table1, 'left', 9)
    set_font_size_column(table1, column_number = 8, size = 8)
    set_header_bold(table1, bold = True)

def subset_dataframe_by_province():
    #initialize list
    province_data = pd.DataFrame(columns = ['linelist_region', 'province', 'classification', 'lineage'])
    linelist_df['LINELIST REGION'] = linelist_df['LINELIST REGION'].astype(str)
    #get unique list of linelist_region
    linelist_region = list(set(linelist_df['LINELIST REGION'].dropna().str.upper()))
    # make an region order 
    region_order = ['CAR','1', '2', '3', '4A', '4B', '5', '6', '7', '8', '9', '10', '11', '12', '13', 'BARMM', 'NCR']
    order = {key: i for i, key in enumerate(region_order)}

    #sort the linelist_region
    linelist_region.sort(key=lambda val: order[val])


    str_linelist_df = linelist_df.astype(str)

    #subset data drame using linelist_region
    for region in linelist_region:
        linelist_region_subset = str_linelist_df[str_linelist_df['LINELIST REGION'].str.upper() == region]
        province_list = list(set(linelist_region_subset['PATIENT ADDRESS'].dropna().str.upper()))
        region_str = region

        #subset dataframe using province
        for province in province_list:
            province_subset = linelist_region_subset[linelist_region_subset['PATIENT ADDRESS'].str.upper() == province]
            province_classification = list(province_subset['criteria_for_selection'])
            province_lineage = list(province_subset['LINEAGE'])
            province_classification = [x.upper() for x in province_classification]
            province_lineage = [x.upper() for x in province_lineage]
            province_str = province

            province_data_subset = pd.DataFrame(columns = ['linelist_region', 'province', 'classification', 'lineage'])
            province_data_subset['province'] = [province_str]
            province_data_subset['linelist_region'] = [region_str]
            province_data_subset['classification'] = [province_classification]
            province_data_subset['lineage'] = [province_lineage]

            province_data = pd.concat([province_data, province_data_subset])

    return province_data

def generate_table2_dataframe(lineage_list,classification_list,province_data):
    lineage_list = [x.upper() for x in lineage_list]

    table2_df = pd.DataFrame(columns=['REPORTNG REGION','REGION-PROVINCE','NO. OF SAMPLES','CLASSIFICATION'] + lineage_list + ['LINEAGE NOT ASSIGNED'])

    temp_df = pd.DataFrame(columns=['REPORTNG REGION','REGION-PROVINCE','NO. OF SAMPLES','CLASSIFICATION'] + lineage_list + ['LINEAGE NOT ASSIGNED']) 

    classification_list = [x.upper() for x in classification_list]

    criteria_dict = {'CLUSTER': 0, 'ADMITTED': 1, 'REINFECTION': 2, 'ROF': 3, 'UNKNOWN EXPOSURE': 4, 'ADMITTED*': 1}


    for index, row in province_data.iterrows():
        temp_df['CLASSIFICATION'] = classification_list
        temp_df['REPORTNG REGION'] = row['linelist_region']
        temp_df['NO. OF SAMPLES'] = 0
        temp_df[lineage_list] = 0
        temp_df['LINEAGE NOT ASSIGNED'] = 0
        temp_df['REGION-PROVINCE'] = row['province']
    
        for i in classification_list:
            for j in row['classification']:
                if (j == i) or (i in j):
                    temp_df.at[criteria_dict[i], 'NO. OF SAMPLES'] += 1

        for x in row['lineage']:
            if (x == ''):
                    classification_lineage_index = row['classification'][row['lineage'].index(x)]
                    
                    temp_df.at[criteria_dict[classification_lineage_index], 'LINEAGE NOT ASSIGNED'] += 1

        lineage_counter = 0

        for x in row['lineage']:
            if x in lineage_list:
                classification_lineage_index = row['classification'][lineage_counter]
                temp_df.at[criteria_dict[classification_lineage_index], x] += 1
          
            lineage_counter += 1

        table2_df = pd.concat([table2_df, temp_df])
    return table2_df

def mark_classification_with_all_zero(table2_df):
    for index, row in table2_df.iterrows():
        if row['NO. OF SAMPLES'] == 0:
            table2_df.at[index, 'CLASSIFICATION'] = row['CLASSIFICATION'] + '*'

    return table2_df

def set_font_color_condition(table):
    for row in table.rows:
        for cell in row.cells:
            if cell.text == '0':
                paragraph = cell.paragraphs[0]
                font = paragraph.runs[0].font
                color = font.color
                color.rgb = RGBColor(0x80, 0x80, 0x80)

def set_classification_column_gray():
    for cell in table2.columns[3].cells:
        if '*' in cell.text:
            cell.text = cell.text.replace('*', '')
            paragraph = cell.paragraphs[0]
            font = paragraph.runs[0].font
            color = font.color
            color.rgb = RGBColor(0x80, 0x80, 0x80) 

def merge_same_text_in_column(column):
    prev_cell = docx.table._Cell
    cell_str = ''
    for cell in table2.columns[column].cells:
        cell_str = prev_cell.text
        if cell_str == cell.text:
            cell.merge(prev_cell).text = cell_str
            cell.text = cell_str
        prev_cell = cell

def set_cell_height(table):
    row_num = 0
    for row in table.rows:
        if row_num == 0:
            pass
        else:
            row.height = Inches(0.71)
        row_num += 1


if __name__ == '__main__':
    print('Generating Linelist...')
    pd.options.mode.chained_assignment = None  # default='warn'
    # open a document
    document = Document("linelist.docx")

    add_line_break(document)

    make_paragraph(document, 'Calibri', 14, WD_ALIGN_PARAGRAPH.CENTER, 'RESULT LINELIST FOR SARS-COV-2 WHOLE GENOME SEQUENCING (GECO-PH)', bold = True) #title
    make_paragraph(document, 'Calibri', 10.5, WD_ALIGN_PARAGRAPH.CENTER, date.today().strftime('%d %B %Y'), False)  #add date today

    add_line_break(document)

    #read csv file from redcap report
    linelist_df = pd.read_csv('linelist_EB_06232022.csv')
    linelist_df.index = linelist_df.index + 1
    linelist_df.rename(columns = {'lab_id':'RITM Lab ID', 'uic':'UIC', 'age': 'AGE', 'sex': 'SEX','city_municipality': 'PATIENT ADDRESS','sending_facility':'SENDING FACILITY', 'date_specimen_collection':'DATE OF COLLECTION', 'pango_lineage': 'LINEAGE', 'linelist_region':'LINELIST REGION'}, inplace = True)

    # put No. column using index
    linelist_df['NO'] = linelist_df.index

    linelist_df = add_cov22(linelist_df)

    linelist_df = make_outcome_column(linelist_df)

    linelist_df = combine_lineage(linelist_df)

    #add remarks column to dataframe
    linelist_df['REMARKS'] = np.nan
    linelist_df['REMARKS'] = ''

    date_list = sort_and_manipulate_date_columns()

    appendix_df = pd.DataFrame(columns=['Date specimens received at RITM'])
    date_received = []
    collection_period = []
    collection_period_subset = []
    num_samples_received = []
    rejected_samples = []
    num_rejected_samples = []
    num_lineage_assigned = []
    num_poor_genome = []
    lineage_list = []
    unique_lineage_list = []
    key_points = []

    subset_and_generate_linelist_table()

    populate_appendix_df()

    document.add_page_break()
    make_paragraph(document, 'Calibri', 10, WD_ALIGN_PARAGRAPH.CENTER, 'APPENDIX: RUN SUMMARY', bold = True)
    add_line_break(document)
    make_paragraph(document, 'Calibri', 10, WD_ALIGN_PARAGRAPH.LEFT, 'Table 1. Summary of sample processing details and key points. ', bold = False)

    # Generate table in for appendix_df in docx file 

    table1 = document.add_table(rows=appendix_df.shape[0]+1, cols=appendix_df.shape[1])
    table1.style = 'Table Grid'

    generate_table(appendix_df, table1)

    # center text in cells
    set_alignment(table1, 'center')

    color_header(table1, '084C22')

    set_table_1_params()

    # TABLE 2
    document.add_page_break()
    make_paragraph(document, 'Calibri', 10, WD_ALIGN_PARAGRAPH.LEFT, 'Table 2. Breakdown of results for EB-endorsed samples processed per region/province.', bold = False)


    breakdown_df = pd.DataFrame(columns=['REPORTNG REGION'])
    #assign all unique regions as the value of the column
    breakdown_df['REGION-PROVINCE'] = list(set(linelist_df['PATIENT ADDRESS'].dropna()))

    province_list = list(breakdown_df['REGION-PROVINCE'])


    classification_list = ['Cluster', 'Admitted', 'Reinfection', 'ROF', 'Unknown Exposure']

    lineage_list = list(set(linelist_df['LINEAGE'][linelist_df['LINEAGE'] != '']))
    test_df = pd.DataFrame(columns=['REPORTNG REGION','REGION-PROVINCE','NO. OF SAMPLES','CLASSIFICATION'] + lineage_list + ['LINEAGE NOT ASSIGNED'])

    classification_table = []

    for place in province_list:
        for x in classification_list:
            classification_table.append(x)


    province_data = subset_dataframe_by_province()

    table2_df = generate_table2_dataframe(lineage_list,classification_list,province_data)

    table2_df = table2_df.reset_index(drop = True)

    table2_df = mark_classification_with_all_zero(table2_df)

    #rename lineage columns add 'Lineage Assigned:'
    for x in lineage_list:
        table2_df.rename(columns = {x: 'Lineage Assigned: ' + x}, inplace = True)

    table2 = document.add_table(rows=table2_df.shape[0]+1, cols=table2_df.shape[1])
    table2.style = 'Table Grid'

    generate_table(table2_df, table2)

    color_header(table2, '084C22')
    set_font_size(table2, 9)
    set_header_bold(table2, True)
    set_alignment(table2, 'center')
    set_font_size(table2, 9)

    set_font_color_condition(table2)

    set_classification_column_gray()

    merge_same_text_in_column(0)
    merge_same_text_in_column(1)
    set_alignment_column(table2,'center',0)
    set_alignment_column(table2,'center',1)
    set_alignment_column(table2,'center',3)

    # save the doc
    document.save('test.docx')






